from fastapi import APIRouter
from fastapi.responses import FileResponse
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.style import WD_STYLE_TYPE
import os
from datetime import datetime

router = APIRouter(tags=["Documentation"])

# Alias for backward compatibility
@router.get("/api/architecture/download")
async def download_architecture_alias():
    """Alias endpoint for architecture download"""
    return await download_architecture_document()

@router.get("/api/docs/architecture-word")
async def download_architecture_document():
    """Generate and download Architecture Overview as Word document"""
    
    doc = Document()
    
    # Set up styles
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Calibri'
    font.size = Pt(11)
    
    # Title
    title = doc.add_heading('INFUSE-AI PLATFORM', 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    subtitle = doc.add_paragraph('Architecture & Modularity Overview | v2.0.0')
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    doc.add_paragraph(f'Generated: {datetime.now().strftime("%Y-%m-%d %H:%M")}')
    doc.add_paragraph()
    
    # System Overview Stats
    doc.add_heading('SYSTEM OVERVIEW', level=1)
    
    stats_table = doc.add_table(rows=1, cols=3)
    stats_table.style = 'Table Grid'
    stats_table.alignment = WD_TABLE_ALIGNMENT.CENTER
    
    cells = stats_table.rows[0].cells
    cells[0].text = '67\nAPI Routes'
    cells[1].text = '72\nBackend Modules'
    cells[2].text = '60+\nDB Collections'
    
    for cell in cells:
        cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    doc.add_paragraph()
    
    # Architecture Diagram
    doc.add_heading('SYSTEM ARCHITECTURE', level=1)
    
    arch_text = """
┌─────────────────────────────────────────────────────────────────────────────────┐
│                              FRONTEND (React + Vite)                             │
├─────────────────────────────────────────────────────────────────────────────────┤
│  PUBLIC PAGES    │  DASHBOARDS      │  ADMIN           │  COMPONENTS            │
│  Landing         │  Doctor          │  Unified         │  HealthTrack (14)      │
│  Login/Register  │  Patient         │  Enterprise      │  SecureSphere (13)     │
│  Downloads       │  HealthTrack     │  Super Admin     │  Dashboard (15)        │
│  Developer       │  SecureSphere    │  Config          │  UI/shadcn (40+)       │
└─────────────────────────────────────────────────────────────────────────────────┘
                                        │
                                        ▼
┌─────────────────────────────────────────────────────────────────────────────────┐
│                            BACKEND (FastAPI + Python)                            │
├─────────────────────────────────────────────────────────────────────────────────┤
│                              API ROUTER (/api)                                   │
│         │                    │                    │                    │         │
│    CORE            HEALTHTRACK PRO      SECURESPHERE          ENTERPRISE        │
│  auth              patients             url_scanner           sso               │
│  organization      appointments         sms_analyzer          bulk_export       │
│  subscription      vitals               threat_score          ip_whitelist      │
│  billing           prescriptions        devices               audit_logs        │
│  dashboard         lab_tests            iot_security          api_keys          │
│  language          ai_analysis          automotive            licenses          │
│  notifications     video_consent        gsm_fraud             analytics         │
│                    telemedicine         telecom                                  │
│                    abdm                 vran                                     │
│                                                                                  │
│  INTEGRATIONS: webhooks | developer_portal | oem_sdk | feature_flags | alerts   │
└─────────────────────────────────────────────────────────────────────────────────┘
                                        │
                                        ▼
┌─────────────────────────────────────────────────────────────────────────────────┐
│                              DATABASE (MongoDB)                                  │
│  users | doctors | patients | organizations | sessions | appointments           │
│  prescriptions | vitals | health_records | video_consents | drug_database       │
│  security_devices | url_scans | sms_analyses | threat_alerts | threat_scores    │
│  api_keys | audit_logs | webhooks | subscriptions | billing_records             │
└─────────────────────────────────────────────────────────────────────────────────┘
"""
    
    arch_para = doc.add_paragraph()
    arch_run = arch_para.add_run(arch_text)
    arch_run.font.name = 'Courier New'
    arch_run.font.size = Pt(8)
    
    doc.add_paragraph()
    
    # Two Products Section
    doc.add_heading('PRODUCT MODULES', level=1)
    
    # HealthTrack Pro
    doc.add_heading('HealthTrack Pro (Healthcare Platform)', level=2)
    ht_table = doc.add_table(rows=6, cols=2)
    ht_table.style = 'Table Grid'
    
    ht_data = [
        ('Routes', '/api/healthtrack/*'),
        ('API Endpoints', '12 endpoints'),
        ('React Components', '14 components'),
        ('DB Collections', '11 collections'),
        ('User Roles', 'Admin, Doctor, Patient, Enterprise'),
        ('Features', 'Patients, Appointments, Vitals, Prescriptions, Lab Tests, AI Analysis, Video Consent, Telemedicine, ABDM Integration'),
    ]
    
    for i, (label, value) in enumerate(ht_data):
        ht_table.rows[i].cells[0].text = label
        ht_table.rows[i].cells[1].text = value
    
    doc.add_paragraph()
    
    # SecureSphere
    doc.add_heading('SecureSphere (Cybersecurity Platform)', level=2)
    ss_table = doc.add_table(rows=6, cols=2)
    ss_table.style = 'Table Grid'
    
    ss_data = [
        ('Routes', '/api/securesphere/*'),
        ('API Endpoints', '11 endpoints'),
        ('React Components', '13 components'),
        ('DB Collections', '20 collections'),
        ('Tier Levels', 'Consumer, Enterprise, Telecom, Automotive'),
        ('Features', 'URL Scanner, SMS Analyzer, Threat Scoring, Device Registry, IoT Security, Automotive Security, GSM Fraud Detection, vRAN Integration'),
    ]
    
    for i, (label, value) in enumerate(ss_data):
        ss_table.rows[i].cells[0].text = label
        ss_table.rows[i].cells[1].text = value
    
    doc.add_paragraph()
    
    # Module Structure
    doc.add_heading('MODULE STRUCTURE', level=1)
    
    doc.add_heading('Backend Structure', level=2)
    backend_structure = """
/app/backend/
├── server.py              # Main FastAPI application
├── dependencies.py        # Shared dependencies
├── routes/
│   ├── auth.py            # Authentication & JWT
│   ├── patients.py        # Patient CRUD operations
│   ├── appointments.py    # Appointment scheduling
│   ├── vitals.py          # Vitals recording
│   ├── ai_health_analysis.py  # AI-powered health insights
│   ├── video_consent.py   # Video consent management
│   ├── securesphere/
│   │   ├── url_scanner.py
│   │   ├── sms_analyzer.py
│   │   ├── threat_scoring.py
│   │   ├── iot_security.py
│   │   └── automotive_security.py
│   ├── enterprise/
│   │   ├── sso.py
│   │   ├── bulk_export.py
│   │   └── audit_logs.py
│   └── admin/
│       ├── api_keys.py
│       └── license_management.py
└── services/
    └── alerts_service.py
"""
    backend_para = doc.add_paragraph()
    backend_run = backend_para.add_run(backend_structure)
    backend_run.font.name = 'Courier New'
    backend_run.font.size = Pt(9)
    
    doc.add_heading('Frontend Structure', level=2)
    frontend_structure = """
/app/frontend/src/
├── App.js                 # Router configuration
├── pages/
│   ├── Login.jsx
│   ├── Register.jsx
│   ├── dashboard/
│   │   ├── DoctorDashboard.jsx
│   │   ├── PatientDashboard.jsx
│   │   ├── HealthTrackPro.jsx
│   │   └── SecureSphere.jsx
│   ├── admin/
│   │   ├── UnifiedAdminDashboard.jsx
│   │   ├── EnterpriseAdminDashboard.jsx
│   │   └── SuperAdminPanel.jsx
│   └── securesphere/
│       ├── SecureSphereDashboard.jsx
│       └── ThreatCenter.jsx
├── components/
│   ├── healthtrack/       # 14 components
│   ├── securesphere/      # 13 components
│   ├── dashboard/         # 15 components
│   └── ui/                # shadcn (40+ components)
└── contexts/
    ├── AuthContext.jsx
    └── LanguageContext.jsx
"""
    frontend_para = doc.add_paragraph()
    frontend_run = frontend_para.add_run(frontend_structure)
    frontend_run.font.name = 'Courier New'
    frontend_run.font.size = Pt(9)
    
    doc.add_paragraph()
    
    # API Endpoints
    doc.add_heading('KEY API ENDPOINTS', level=1)
    
    api_table = doc.add_table(rows=7, cols=3)
    api_table.style = 'Table Grid'
    
    # Headers
    api_table.rows[0].cells[0].text = 'CORE'
    api_table.rows[0].cells[1].text = 'HEALTHTRACK PRO'
    api_table.rows[0].cells[2].text = 'SECURESPHERE'
    
    core_apis = ['/api/auth/login', '/api/auth/register', '/api/organization', '/api/subscription', '/api/dashboard', '/api/language']
    ht_apis = ['/api/healthtrack/patients', '/api/healthtrack/appointments', '/api/healthtrack/vitals', '/api/healthtrack/prescriptions', '/api/healthtrack/ai-analysis', '/api/healthtrack/video-consent']
    ss_apis = ['/api/securesphere/url-scanner', '/api/securesphere/sms-analyzer', '/api/securesphere/threat-scoring', '/api/securesphere/devices', '/api/securesphere/iot', '/api/securesphere/automotive']
    
    for i in range(6):
        api_table.rows[i+1].cells[0].text = core_apis[i]
        api_table.rows[i+1].cells[1].text = ht_apis[i]
        api_table.rows[i+1].cells[2].text = ss_apis[i]
    
    doc.add_paragraph()
    
    # Tech Stack
    doc.add_heading('TECHNOLOGY STACK', level=1)
    
    tech_table = doc.add_table(rows=7, cols=2)
    tech_table.style = 'Table Grid'
    
    tech_data = [
        ('Frontend', 'React 18, Vite, TailwindCSS, shadcn/ui'),
        ('Backend', 'FastAPI, Python 3.11, Pydantic'),
        ('Database', 'MongoDB Atlas, Motor (async driver)'),
        ('Authentication', 'JWT, bcrypt, OAuth2'),
        ('Charts', 'Recharts'),
        ('Mobile', 'Capacitor (iOS/Android)'),
        ('Internationalization', '6 languages supported'),
    ]
    
    for i, (label, value) in enumerate(tech_data):
        tech_table.rows[i].cells[0].text = label
        tech_table.rows[i].cells[1].text = value
    
    doc.add_paragraph()
    
    # Third-party Integrations
    doc.add_heading('THIRD-PARTY INTEGRATIONS', level=1)
    
    int_table = doc.add_table(rows=7, cols=3)
    int_table.style = 'Table Grid'
    
    int_table.rows[0].cells[0].text = 'Service'
    int_table.rows[0].cells[1].text = 'Purpose'
    int_table.rows[0].cells[2].text = 'Status'
    
    integrations = [
        ('OpenAI/Gemini', 'AI Health Analysis', 'ACTIVE'),
        ('Razorpay', 'Payment Processing', 'ACTIVE'),
        ('Twilio', 'SMS/WhatsApp Alerts', 'MOCKED'),
        ('SendGrid', 'Email Notifications', 'MOCKED'),
        ('Recharts', 'Dashboard Charts', 'ACTIVE'),
        ('Capacitor', 'Mobile App', 'ACTIVE'),
    ]
    
    for i, (service, purpose, status) in enumerate(integrations):
        int_table.rows[i+1].cells[0].text = service
        int_table.rows[i+1].cells[1].text = purpose
        int_table.rows[i+1].cells[2].text = status
    
    doc.add_paragraph()
    
    # Feature Modularity
    doc.add_heading('FEATURE MODULARITY', level=1)
    
    feat_table = doc.add_table(rows=13, cols=3)
    feat_table.style = 'Table Grid'
    
    feat_table.rows[0].cells[0].text = 'Feature'
    feat_table.rows[0].cells[1].text = 'Module File'
    feat_table.rows[0].cells[2].text = 'Status'
    
    features = [
        ('AI Health Analysis', 'ai_health_analysis.py', 'ON'),
        ('Video Consent', 'video_consent.py', 'ON'),
        ('ABDM Integration', 'abdm.py', 'ON'),
        ('Wearable Sync', 'wearable_devices.py', 'ON'),
        ('URL Scanner', 'url_scanner.py', 'ON'),
        ('SMS Analyzer', 'sms_analyzer.py', 'ON'),
        ('IoT Security', 'iot_security.py', 'ON'),
        ('Automotive Security', 'automotive_security.py', 'ON'),
        ('vRAN Integration', 'vran_api.py', 'ON'),
        ('SSO Authentication', 'sso.py', 'ON'),
        ('Webhooks', 'webhooks.py', 'ON'),
        ('Real Alerts', 'alerts.py', 'MOCKED'),
    ]
    
    for i, (name, module, status) in enumerate(features):
        feat_table.rows[i+1].cells[0].text = name
        feat_table.rows[i+1].cells[1].text = module
        feat_table.rows[i+1].cells[2].text = status
    
    doc.add_paragraph()
    
    # Test Credentials
    doc.add_heading('TEST CREDENTIALS', level=1)
    
    cred_table = doc.add_table(rows=5, cols=3)
    cred_table.style = 'Table Grid'
    
    cred_table.rows[0].cells[0].text = 'Role'
    cred_table.rows[0].cells[1].text = 'Email'
    cred_table.rows[0].cells[2].text = 'Password'
    
    credentials = [
        ('Admin', 'admin@infuse.demo', 'admin1234'),
        ('Doctor', 'doctor.priya@infuse.demo', 'demo1234'),
        ('Enterprise', 'enterprise@infuse.demo', 'demo1234'),
        ('Patient', 'patient.kumar@infuse.demo', 'demo1234'),
    ]
    
    for i, (role, email, password) in enumerate(credentials):
        cred_table.rows[i+1].cells[0].text = role
        cred_table.rows[i+1].cells[1].text = email
        cred_table.rows[i+1].cells[2].text = password
    
    doc.add_paragraph()
    
    # Master URL
    url_para = doc.add_paragraph()
    url_para.add_run('MASTER URL: ').bold = True
    url_para.add_run('https://qa-track-suite.preview.emergentagent.com')
    
    doc.add_paragraph()
    
    # Footer
    footer = doc.add_paragraph()
    footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    footer.add_run('─' * 80)
    doc.add_paragraph('Infuse-AI Platform Architecture Document').alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # Save document
    file_path = '/app/frontend/public/downloads/Infuse-AI-Architecture-Overview.docx'
    os.makedirs(os.path.dirname(file_path), exist_ok=True)
    doc.save(file_path)
    
    return FileResponse(
        path=file_path,
        filename='Infuse-AI-Architecture-Overview.docx',
        media_type='application/vnd.openxmlformats-officedocument.wordprocessingml.document'
    )
